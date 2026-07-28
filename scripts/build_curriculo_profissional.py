from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    HRFlowable,
    KeepTogether,
    PageTemplate,
    Paragraph,
    Spacer,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOC = ROOT / "output" / "doc" / "Curriculo_Gabriel_Schwingel_Conci.docx"
OUTPUT_PDF = ROOT / "files" / "Gabriel_Conci.pdf"

FONT = "Arial"
INK = RGBColor(31, 42, 55)
MUTED = RGBColor(71, 85, 105)
ACCENT = RGBColor(13, 112, 82)
LINK = RGBColor(5, 92, 130)

PDF_INK = colors.HexColor("#1F2A37")
PDF_MUTED = colors.HexColor("#475569")
PDF_ACCENT = colors.HexColor("#0D7052")
PDF_LINK = colors.HexColor("#055C82")


def set_cell_margins(section) -> None:
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)


def set_run_font(run, *, size: float, bold: bool = False, color=INK) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def set_keep_with_next(paragraph, enabled: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = enabled


def set_keep_together(paragraph, enabled: bool = True) -> None:
    paragraph.paragraph_format.keep_together = enabled


def add_bottom_border(paragraph, *, color: str = "0D7052", size: str = "10") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text: str, url: str):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_fonts = OxmlElement("w:rFonts")
    run_fonts.set(qn("w:ascii"), FONT)
    run_fonts.set(qn("w:hAnsi"), FONT)
    run_fonts.set(qn("w:eastAsia"), FONT)
    run_properties.append(run_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "055C82")
    run_properties.append(color)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    run_properties.append(size)

    size_complex = OxmlElement("w:szCs")
    size_complex.set(qn("w:val"), "18")
    run_properties.append(size_complex)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)
    node_text = OxmlElement("w:t")
    node_text.text = text
    run.append(node_text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "Resume Bullet" not in [style.name for style in document.styles]:
        bullet = document.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = document.styles["Resume Bullet"]
    bullet.base_style = document.styles["Normal"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    bullet.font.size = Pt(9.8)
    bullet.font.color.rgb = INK
    bullet.paragraph_format.left_indent = Cm(0.42)
    bullet.paragraph_format.first_line_indent = Cm(-0.29)
    bullet.paragraph_format.right_indent = Cm(0)
    bullet.paragraph_format.space_after = Pt(1.4)
    bullet.paragraph_format.line_spacing = 1.0


def add_header(document: Document) -> None:
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    run = name.add_run("GABRIEL SCHWINGEL CONCI")
    set_run_font(run, size=20.5, bold=True, color=INK)
    set_keep_with_next(name)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run(
        "Desenvolvedor Full Stack | Python, FastAPI, Automação e IA Aplicada"
    )
    set_run_font(run, size=10.6, bold=True, color=MUTED)
    set_keep_with_next(title)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(1.5)
    set_run_font(
        contact.add_run("Gravataí, RS  |  +55 (51) 99446-4827  |  "),
        size=9.1,
        color=MUTED,
    )
    add_hyperlink(contact, "gabrielsconci@gmail.com", "mailto:gabrielsconci@gmail.com")
    set_keep_with_next(contact)

    links = document.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.paragraph_format.space_after = Pt(5)
    add_hyperlink(
        links,
        "linkedin.com/in/gabriel-schwingel-conci-a0528a344",
        "https://www.linkedin.com/in/gabriel-schwingel-conci-a0528a344/",
    )
    set_run_font(links.add_run("  |  "), size=9, color=MUTED)
    add_hyperlink(links, "github.com/3Gb3", "https://github.com/3Gb3")
    set_run_font(links.add_run("  |  "), size=9, color=MUTED)
    add_hyperlink(links, "3gb3.github.io/3Gb3", "https://3gb3.github.io/3Gb3/")
    set_keep_with_next(links)


def add_section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(2.6)
    run = paragraph.add_run(text.upper())
    set_run_font(run, size=10.1, bold=True, color=ACCENT)
    add_bottom_border(paragraph)
    set_keep_with_next(paragraph)


def add_body_paragraph(document: Document, text: str, *, after: float = 1.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.02
    run = paragraph.add_run(text)
    set_run_font(run, size=10)
    set_keep_together(paragraph)


def add_labeled_line(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1.2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run_font(paragraph.add_run(f"{label}: "), size=9.8, bold=True, color=INK)
    set_run_font(paragraph.add_run(value), size=9.8, color=INK)
    set_keep_together(paragraph)


def add_role_heading(document: Document, company: str, role: str, dates: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2.3)
    paragraph.paragraph_format.space_after = Pt(1.4)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Cm(17.8), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run_font(paragraph.add_run(company), size=10, bold=True, color=INK)
    set_run_font(paragraph.add_run(f" | {role}"), size=10, bold=True, color=MUTED)
    set_run_font(paragraph.add_run(f"\t{dates}"), size=9.6, bold=True, color=MUTED)
    set_keep_with_next(paragraph)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Resume Bullet")
    set_run_font(paragraph.add_run("- "), size=9.8, color=ACCENT)
    set_run_font(paragraph.add_run(text), size=9.8, color=INK)
    set_keep_together(paragraph)


def add_project(document: Document, name: str, tech: str, description: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.42)
    paragraph.paragraph_format.first_line_indent = Cm(-0.29)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run_font(paragraph.add_run("- "), size=9.7, color=ACCENT)
    set_run_font(paragraph.add_run(name), size=9.7, bold=True, color=INK)
    set_run_font(paragraph.add_run(f" | {tech} - "), size=9.7, bold=True, color=MUTED)
    set_run_font(paragraph.add_run(description), size=9.7, color=INK)
    set_keep_together(paragraph)


def build_document() -> Document:
    document = Document()
    if len(document.sections) > 1:
        while len(document.sections) > 1:
            document.sections[-1].start_type = WD_SECTION.CONTINUOUS

    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    set_cell_margins(section)
    configure_styles(document)

    properties = document.core_properties
    properties.title = "Currículo - Gabriel Schwingel Conci"
    properties.author = "Gabriel Schwingel Conci"
    properties.subject = (
        "Desenvolvedor Full Stack | Python, FastAPI, Automação e IA Aplicada"
    )
    properties.keywords = (
        "Python, FastAPI, automação, chatbots, inteligência artificial, "
        "backend, PostgreSQL, RabbitMQ, Go"
    )
    properties.comments = "Currículo profissional em formato ATS, uma coluna."

    add_header(document)

    add_section_heading(document, "Resumo profissional")
    add_body_paragraph(
        document,
        "Desenvolvedor Full Stack em formação, com atuação prática na Lojas "
        "Quero-Quero em APIs, automações, chatbots e soluções de inteligência "
        "artificial. Experiência com desenvolvimento backend, integrações, "
        "mensageria e qualidade de fluxos conversacionais.",
    )

    add_section_heading(document, "Competências técnicas")
    add_labeled_line(document, "Linguagens", "Python, JavaScript, Go, SQL, HTML5 e CSS3")
    add_labeled_line(
        document,
        "Backend e IA",
        "FastAPI, Flask, Node.js, ChatGraph, OpenRouter, OpenAI, Ollama, APIs REST e tool calling",
    )
    add_labeled_line(
        document,
        "Dados e mensageria",
        "PostgreSQL, Firebase/Firestore, RabbitMQ, Pandas e DBeaver",
    )
    add_labeled_line(
        document,
        "Arquitetura e entrega",
        "Arquitetura hexagonal, três camadas, Docker, Git/GitHub, Postman e WSL",
    )

    add_section_heading(document, "Experiência profissional")
    add_role_heading(
        document,
        "LOJAS QUERO-QUERO",
        "Estagiário de Desenvolvimento | Business Tech / QQTech",
        "mai. 2026 - atual",
    )
    add_bullet(
        document,
        "Desenvolvo APIs, automações e integrações com Python, FastAPI, "
        "JavaScript/Node.js, Go, PostgreSQL e RabbitMQ.",
    )
    add_bullet(
        document,
        "Implementei chatbot de atendimento para Instagram e Facebook, com "
        "interpretação de intenção, chamadas de ferramentas, roteamento de "
        "fluxos e respostas automatizadas.",
    )
    add_bullet(
        document,
        "Modernizei fluxos legados do principal canal de atendimento, "
        "reorganizando menus e direcionamentos para facilitar a evolução da solução.",
    )
    add_bullet(
        document,
        "Co-desenvolvo um Memory Service para armazenamento e recuperação de "
        "contexto em soluções internas de IA.",
    )
    add_bullet(
        document,
        "Desenvolvo o sistema interno Oferte e Ganhe para controle de talões, "
        "com APIs, relatórios e recursos de inteligência artificial integrados.",
    )

    add_role_heading(
        document,
        "LOJAS QUERO-QUERO",
        "Jovem Aprendiz | QA de Chatbots",
        "nov. 2025 - mai. 2026",
    )
    add_bullet(
        document,
        "Executei testes funcionais de chatbots e respostas de IA, validando "
        "fluxos conversacionais, regras de negócio, exceções e experiência do usuário.",
    )
    add_bullet(
        document,
        "Mapeei fluxos no Miro, reproduzi e documentei falhas com evidências e "
        "acompanhei correções até a validação final.",
    )

    add_section_heading(document, "Projetos em destaque")
    add_project(
        document,
        "Ventude Planner",
        "JavaScript, Firebase e Chart.js",
        "planner pessoal com autenticação, rotinas, metas, persistência e dashboards responsivos.",
    )
    add_project(
        document,
        "Code Logic",
        "Python, Flask, Firebase e IA",
        "plataforma educacional com módulos, exercícios e feedback personalizado.",
    )
    add_project(
        document,
        "Projeto BNCC",
        "Python, Tkinter e processamento de PDF",
        "aplicação desktop que identifica códigos da BNCC e apoia o planejamento pedagógico.",
    )

    add_section_heading(document, "Formação e cursos")
    education = document.add_paragraph()
    education.paragraph_format.space_after = Pt(1.5)
    education.paragraph_format.tab_stops.add_tab_stop(
        Cm(17.8), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run_font(
        education.add_run("CESUCA | Tecnologia em Análise e Desenvolvimento de Sistemas"),
        size=9.55,
        bold=True,
        color=INK,
    )
    set_run_font(
        education.add_run("\t2025 - atual | 4º semestre"),
        size=9.2,
        bold=True,
        color=MUTED,
    )
    set_keep_with_next(education)
    add_labeled_line(
        document,
        "Formação complementar",
        "Python do Básico ao Avançado (Udemy, em andamento) e Lógica da Programação em VisualG e Python (Udemy, 2025)",
    )

    return document


class RoleLine(Flowable):
    def __init__(self, company: str, role: str, dates: str, width: float):
        super().__init__()
        self.company = company
        self.role = role
        self.dates = dates
        self.width = width
        self.height = 13

    def draw(self) -> None:
        canvas = self.canv
        canvas.saveState()
        canvas.setFillColor(PDF_INK)
        canvas.setFont("Arial-Bold", 9.7)
        canvas.drawString(0, 2.5, self.company)
        company_width = pdfmetrics.stringWidth(self.company, "Arial-Bold", 9.7)
        canvas.setFillColor(PDF_MUTED)
        canvas.drawString(company_width, 2.5, f" | {self.role}")
        canvas.setFont("Arial-Bold", 9.2)
        canvas.drawRightString(self.width, 2.5, self.dates)
        canvas.restoreState()


class EducationLine(Flowable):
    def __init__(self, text: str, dates: str, width: float):
        super().__init__()
        self.text = text
        self.dates = dates
        self.width = width
        self.height = 13

    def draw(self) -> None:
        canvas = self.canv
        canvas.saveState()
        canvas.setFillColor(PDF_INK)
        canvas.setFont("Arial-Bold", 9.6)
        canvas.drawString(0, 2.5, self.text)
        canvas.setFillColor(PDF_MUTED)
        canvas.setFont("Arial-Bold", 9.1)
        canvas.drawRightString(self.width, 2.5, self.dates)
        canvas.restoreState()


def register_pdf_fonts() -> None:
    font_dir = Path("C:/Windows/Fonts")
    font_files = {
        "Arial": font_dir / "arial.ttf",
        "Arial-Bold": font_dir / "arialbd.ttf",
        "Arial-Italic": font_dir / "ariali.ttf",
        "Arial-BoldItalic": font_dir / "arialbi.ttf",
    }
    for name, path in font_files.items():
        pdfmetrics.registerFont(TTFont(name, str(path)))
    pdfmetrics.registerFontFamily(
        "Arial",
        normal="Arial",
        bold="Arial-Bold",
        italic="Arial-Italic",
        boldItalic="Arial-BoldItalic",
    )


def pdf_styles() -> dict[str, ParagraphStyle]:
    return {
        "name": ParagraphStyle(
            "Name",
            fontName="Arial-Bold",
            fontSize=20.5,
            leading=22,
            textColor=PDF_INK,
            alignment=TA_CENTER,
            spaceAfter=1.5,
        ),
        "title": ParagraphStyle(
            "Title",
            fontName="Arial-Bold",
            fontSize=10.6,
            leading=12,
            textColor=PDF_MUTED,
            alignment=TA_CENTER,
            spaceAfter=3,
        ),
        "contact": ParagraphStyle(
            "Contact",
            fontName="Arial",
            fontSize=9.2,
            leading=11,
            textColor=PDF_MUTED,
            alignment=TA_CENTER,
            spaceAfter=1,
        ),
        "links": ParagraphStyle(
            "Links",
            fontName="Arial",
            fontSize=9,
            leading=11,
            textColor=PDF_LINK,
            alignment=TA_CENTER,
            spaceAfter=4,
        ),
        "section": ParagraphStyle(
            "Section",
            fontName="Arial-Bold",
            fontSize=10.2,
            leading=11.8,
            textColor=PDF_ACCENT,
            spaceBefore=4.5,
            spaceAfter=1.2,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "Body",
            fontName="Arial",
            fontSize=10,
            leading=11.7,
            textColor=PDF_INK,
            spaceAfter=1,
        ),
        "skill": ParagraphStyle(
            "Skill",
            fontName="Arial",
            fontSize=9.8,
            leading=11.2,
            textColor=PDF_INK,
            spaceAfter=0.8,
        ),
        "bullet": ParagraphStyle(
            "Bullet",
            fontName="Arial",
            fontSize=9.8,
            leading=11.4,
            textColor=PDF_INK,
            leftIndent=10,
            firstLineIndent=-8,
            bulletIndent=1,
            spaceAfter=1.5,
        ),
        "project": ParagraphStyle(
            "Project",
            fontName="Arial",
            fontSize=9.7,
            leading=11.2,
            textColor=PDF_INK,
            leftIndent=10,
            firstLineIndent=-8,
            bulletIndent=1,
            spaceAfter=1.4,
        ),
    }


def section_story(text: str, styles: dict[str, ParagraphStyle]) -> list:
    return [
        Paragraph(text.upper(), styles["section"]),
        HRFlowable(
            width="100%",
            thickness=0.75,
            color=PDF_ACCENT,
            spaceBefore=0,
            spaceAfter=2,
        ),
    ]


def skill_story(
    label: str, value: str, styles: dict[str, ParagraphStyle]
) -> Paragraph:
    return Paragraph(f"<b>{label}:</b> {value}", styles["skill"])


def bullet_story(text: str, styles: dict[str, ParagraphStyle]) -> Paragraph:
    return Paragraph(f"- {text}", styles["bullet"])


def project_story(
    name: str,
    tech: str,
    description: str,
    styles: dict[str, ParagraphStyle],
) -> Paragraph:
    return Paragraph(
        f"- <b>{name}</b> | <b><font color='#475569'>{tech} -</font></b> {description}",
        styles["project"],
    )


def build_pdf() -> None:
    register_pdf_fonts()
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    width, height = A4
    left = right = 1.7 * cm
    top = 1.25 * cm
    bottom = 1.15 * cm
    frame_width = width - left - right
    frame = Frame(
        left,
        bottom,
        frame_width,
        height - top - bottom,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )

    document = BaseDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        leftMargin=left,
        rightMargin=right,
        topMargin=top,
        bottomMargin=bottom,
        title="Currículo - Gabriel Schwingel Conci",
        author="Gabriel Schwingel Conci",
        subject="Desenvolvedor Full Stack | Python, FastAPI, Automação e IA Aplicada",
        keywords=(
            "Python, FastAPI, automação, chatbots, inteligência artificial, "
            "backend, PostgreSQL, RabbitMQ, Go"
        ),
    )
    document.addPageTemplates([PageTemplate(id="resume", frames=[frame])])
    styles = pdf_styles()

    story: list = [
        Paragraph("GABRIEL SCHWINGEL CONCI", styles["name"]),
        Paragraph(
            "Desenvolvedor Full Stack | Python, FastAPI, Automação e IA Aplicada",
            styles["title"],
        ),
        Paragraph(
            "Gravataí, RS&nbsp;&nbsp;|&nbsp;&nbsp;+55 (51) 99446-4827"
            "&nbsp;&nbsp;|&nbsp;&nbsp;"
            "<link href='mailto:gabrielsconci@gmail.com' color='#055C82'>"
            "gabrielsconci@gmail.com</link>",
            styles["contact"],
        ),
        Paragraph(
            "<link href='https://www.linkedin.com/in/"
            "gabriel-schwingel-conci-a0528a344/' color='#055C82'>"
            "linkedin.com/in/gabriel-schwingel-conci-a0528a344</link>"
            "&nbsp;&nbsp;|&nbsp;&nbsp;"
            "<link href='https://github.com/3Gb3' color='#055C82'>github.com/3Gb3</link>"
            "&nbsp;&nbsp;|&nbsp;&nbsp;"
            "<link href='https://3gb3.github.io/3Gb3/' color='#055C82'>"
            "3gb3.github.io/3Gb3</link>",
            styles["links"],
        ),
    ]

    story += section_story("Resumo profissional", styles)
    story.append(
        Paragraph(
            "Desenvolvedor Full Stack em formação, com atuação prática na Lojas "
            "Quero-Quero em APIs, automações, chatbots e soluções de inteligência "
            "artificial. Experiência com desenvolvimento backend, integrações, "
            "mensageria e qualidade de fluxos conversacionais.",
            styles["body"],
        )
    )

    story += section_story("Competências técnicas", styles)
    story += [
        skill_story(
            "Linguagens", "Python, JavaScript, Go, SQL, HTML5 e CSS3", styles
        ),
        skill_story(
            "Backend e IA",
            "FastAPI, Flask, Node.js, ChatGraph, OpenRouter, OpenAI, Ollama, APIs REST e tool calling",
            styles,
        ),
        skill_story(
            "Dados e mensageria",
            "PostgreSQL, Firebase/Firestore, RabbitMQ, Pandas e DBeaver",
            styles,
        ),
        skill_story(
            "Arquitetura e entrega",
            "Arquitetura hexagonal, três camadas, Docker, Git/GitHub, Postman e WSL",
            styles,
        ),
    ]

    story += section_story("Experiência profissional", styles)
    story.append(
        RoleLine(
            "LOJAS QUERO-QUERO",
            "Estagiário de Desenvolvimento | Business Tech / QQTech",
            "mai. 2026 - atual",
            frame_width,
        )
    )
    story += [
        bullet_story(
            "Desenvolvo APIs, automações e integrações com Python, FastAPI, "
            "JavaScript/Node.js, Go, PostgreSQL e RabbitMQ.",
            styles,
        ),
        bullet_story(
            "Implementei chatbot de atendimento para Instagram e Facebook, com "
            "interpretação de intenção, chamadas de ferramentas, roteamento de "
            "fluxos e respostas automatizadas.",
            styles,
        ),
        bullet_story(
            "Modernizei fluxos legados do principal canal de atendimento, "
            "reorganizando menus e direcionamentos para facilitar a evolução da solução.",
            styles,
        ),
        bullet_story(
            "Co-desenvolvo um Memory Service para armazenamento e recuperação de "
            "contexto em soluções internas de IA.",
            styles,
        ),
        bullet_story(
            "Desenvolvo o sistema interno Oferte e Ganhe para controle de talões, "
            "com APIs, relatórios e recursos de inteligência artificial integrados.",
            styles,
        ),
        RoleLine(
            "LOJAS QUERO-QUERO",
            "Jovem Aprendiz | QA de Chatbots",
            "nov. 2025 - mai. 2026",
            frame_width,
        ),
        bullet_story(
            "Executei testes funcionais de chatbots e respostas de IA, validando "
            "fluxos conversacionais, regras de negócio, exceções e experiência do usuário.",
            styles,
        ),
        bullet_story(
            "Mapeei fluxos no Miro, reproduzi e documentei falhas com evidências e "
            "acompanhei correções até a validação final.",
            styles,
        ),
    ]

    story += section_story("Projetos em destaque", styles)
    story += [
        project_story(
            "Ventude Planner",
            "JavaScript, Firebase e Chart.js",
            "planner pessoal com autenticação, rotinas, metas, persistência e dashboards responsivos.",
            styles,
        ),
        project_story(
            "Code Logic",
            "Python, Flask, Firebase e IA",
            "plataforma educacional com módulos, exercícios e feedback personalizado.",
            styles,
        ),
        project_story(
            "Projeto BNCC",
            "Python, Tkinter e processamento de PDF",
            "aplicação desktop que identifica códigos da BNCC e apoia o planejamento pedagógico.",
            styles,
        ),
    ]

    story += section_story("Formação e cursos", styles)
    story.append(
        KeepTogether(
            [
                EducationLine(
                    "CESUCA | Tecnologia em Análise e Desenvolvimento de Sistemas",
                    "2025 - atual | 4º semestre",
                    frame_width,
                ),
                Paragraph(
                    "<b>Formação complementar:</b> Python do Básico ao Avançado "
                    "(Udemy, em andamento) e Lógica da Programação em VisualG e Python "
                    "(Udemy, 2025)",
                    styles["skill"],
                ),
            ]
        )
    )

    document.build(story)


def main() -> None:
    OUTPUT_DOC.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_DOC)
    build_pdf()
    print(OUTPUT_DOC)
    print(OUTPUT_PDF)


if __name__ == "__main__":
    main()
